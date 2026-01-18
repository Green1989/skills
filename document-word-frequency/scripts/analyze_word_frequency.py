#!/usr/bin/env python3
"""
Document Word Frequency Analyzer

Analyzes word frequency in .doc, .docx, and .pdf documents.
Generates a formatted markdown report with statistics.
"""

import sys
import re
import argparse
from pathlib import Path
from collections import Counter
from typing import Dict, List, Tuple


def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from .docx files using python-docx."""
    try:
        from docx import Document
        doc = Document(docx_path)

        # Collect all text as a single string
        # Note: We iterate through all cells including merged ones to ensure accurate counting
        all_text = ""

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                all_text += text + " "

        # Extract text from tables (all cells, including merged ones)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        all_text += text + " "

        return all_text
    except ImportError:
        print("Error: python-docx is required. Install with: pip install python-docx")
        sys.exit(1)
    except Exception as e:
        print(f"Error reading DOCX file: {e}")
        sys.exit(1)


def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from .pdf files using pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(pdf_path)
        text = ''
        for page in reader.pages:
            text += page.extract_text() + '\n'
        return text
    except ImportError:
        print("Error: pypdf is required. Install with: pip install pypdf")
        sys.exit(1)
    except Exception as e:
        print(f"Error reading PDF file: {e}")
        sys.exit(1)


def extract_text_from_doc(doc_path: str) -> str:
    """Extract text from .doc files using textract."""
    try:
        import textract
        text = textract.process(doc_path).decode('utf-8')
        return text
    except ImportError:
        print("Warning: textract not installed. Trying alternative method...")
        print("Install textract for better .doc support: pip install textract")
        # Fallback: try using antiword if available
        try:
            import subprocess
            result = subprocess.run(['antiword', doc_path],
                                    capture_output=True, text=True)
            if result.returncode == 0:
                return result.stdout
            else:
                print("Error: antiword failed. Please install textract for .doc support.")
                sys.exit(1)
        except FileNotFoundError:
            print("Error: Neither textract nor antiword is available.")
            print("Please install textract: pip install textract")
            sys.exit(1)
    except Exception as e:
        print(f"Error reading DOC file: {e}")
        sys.exit(1)


def extract_text(file_path: str) -> str:
    """Extract text from supported document formats."""
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    extractors = {
        '.docx': extract_text_from_docx,
        '.pdf': extract_text_from_pdf,
        '.doc': extract_text_from_doc,
    }

    if suffix not in extractors:
        raise ValueError(f"Unsupported file format: {suffix}. Supported formats: .doc, .docx, .pdf")

    return extractors[suffix](str(file_path))


def tokenize_text(text: str, language: str = 'mixed') -> List[str]:
    """
    Tokenize text into words.
    Supports Chinese, English, and mixed content.

    Chinese: Use jieba with custom words + regex for comprehensive coverage
    English: Extract complete words

    Args:
        text: The text to tokenize
        language: 'chinese', 'english', or 'mixed' (default)

    Returns:
        List of tokens
    """
    tokens = []

    # For Chinese text, use jieba with custom word additions
    if language in ['chinese', 'mixed']:
        try:
            import jieba

            # Add common multi-character phrases that jieba might split
            # These are commonly searched terms in technical documents
            custom_words = ['提交数据', '响应时间', '检验机构', '检测线', '返回数据']
            for word in custom_words:
                jieba.add_word(word, freq=10000, tag='nz')

            # Use jieba.cut with cut_all=True to find all possible words
            chinese_words = jieba.cut(text, cut_all=True)
            # Filter: only keep words with 2+ Chinese characters
            for word in chinese_words:
                word = word.strip()
                if word:
                    # Check if word has 2+ Chinese characters
                    chinese_chars = re.findall(r'[\u4e00-\u9fff\u3400-\u4dbf\uf900-\ufaff]', word)
                    if len(chinese_chars) >= 2:
                        tokens.append(word)
        except ImportError:
            # Fallback to regex if jieba is not available
            print("Warning: jieba not installed. Using regex fallback. Install with: pip install jieba")
            chinese_pattern = re.compile(r'[\u4e00-\u9fff\u3400-\u4dbf\uf900-\ufaff]{2,}')
            chinese_words = chinese_pattern.findall(text)
            tokens.extend(chinese_words)

    # For English and other languages, extract words
    if language in ['english', 'mixed']:
        # Match ASCII words only (a-z, A-Z, 0-9, underscore, hyphen)
        # Use re.ASCII flag to prevent matching Chinese characters
        words = re.findall(r'\b[a-zA-Z0-9_-]+\b', text, re.ASCII)
        # Filter out pure numbers
        words = [w for w in words if not w.replace('.', '').replace('-', '').replace('_', '').isdigit()]
        tokens.extend(words)

    return tokens


def count_word_frequency(tokens: List[str]) -> Counter:
    """
    Count word frequency from tokens.
    Filters out pure numeric tokens.
    """
    # Filter out tokens that are purely numeric (including decimals)
    filtered_tokens = []
    for token in tokens:
        # Check if token is purely numeric
        test_token = token.replace('.', '').replace('-', '').replace('/', '')
        if not test_token.isdigit() and token:  # Also filter empty strings
            filtered_tokens.append(token)

    return Counter(filtered_tokens)


def generate_markdown_report(
    file_path: str,
    word_counts: Counter,
    total_words: int,
    unique_words: int,
    top_n: int = 100
) -> str:
    """
    Generate a formatted markdown report.

    Args:
        file_path: Path to the analyzed file
        word_counts: Counter with word frequencies
        total_words: Total number of words
        unique_words: Number of unique words
        top_n: Number of top words to display

    Returns:
        Formatted markdown report
    """
    file_name = Path(file_path).name

    # Generate timestamp
    from datetime import datetime
    timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    report = f"""# 文档词频统计报告

## 文档信息

- **文件名**: {file_name}
- **文件路径**: {file_path}
- **总词数**: {total_words:,}
- **不重复词汇数**: {unique_words:,}

---

## 统计摘要

本报告展示了文档中词汇的出现频率统计。统计规则：
- **中文词汇**: 只统计至少2个连续汉字组成的词汇，单个汉字不统计
- **英文单词**: 统计完整的单词（字母数字序列）
- **数字过滤**: 纯数字（包括日期、编号、ID等）已被过滤，不进行统计
- **无过滤**: 未对停用词、标点符号等进行过滤，以反映文档真实情况

---

## 词频统计表 (Top {top_n})

| 排名 | 词汇 | 出现次数 | 频率 (%) |
|------|------|----------|----------|
"""

    # Get top N words
    top_words = word_counts.most_common(top_n)

    for rank, (word, count) in enumerate(top_words, 1):
        frequency = (count / total_words) * 100
        # Escape pipe characters in markdown tables
        display_word = word.replace('|', '\\|')
        report += f"| {rank} | `{display_word}` | {count:,} | {frequency:.2f}% |\n"

    report += f"""

---

## 统计说明

- **统计范围**: 文档中的所有词汇（包括段落和表格）
- **中文处理**: 只统计至少2个连续汉字，单个汉字不统计
- **英文处理**: 按完整单词统计（字母数字序列）
- **数字过滤**: 纯数字（包括日期、编号、ID等）已过滤不统计
- **过滤策略**: 保留停用词、标点符号等，仅过滤纯数字
- **排名依据**: 按出现频次降序排列

---

## 图表展示

### 频率分布

前 {min(20, len(top_words))} 个高频词汇的频率分布：

```mermaid
pie
    title Top {min(20, len(top_words))} 高频词汇分布
"""

    # Add top 20 words to pie chart
    for word, count in top_words[:20]:
        # Clean word for mermaid (remove special characters)
        clean_word = re.sub(r'[^\w\u4e00-\u9fff-]', '_', word)[:20]
        report += f'    "{clean_word}" : {count}\n'

    report += "```\n"

    report += f"""

### 词频柱状图 (Top 20)

| 词汇 | 出现次数 | 可视化 |
|------|----------|--------|
"""

    # Add bar chart for top 20
    max_count = top_words[0][1] if top_words else 0
    for word, count in top_words[:20]:
        bar_length = int((count / max_count) * 50)
        bar = '█' * bar_length
        display_word = word.replace('|', '\\|')
        report += f"| `{display_word}` | {count:,} | {bar} |\n"

    report += f"""

---

**报告生成时间**: {timestamp}

---

*本报告由 Claude Code Document Word Frequency Analyzer 自动生成*
"""

    return report


def main():
    parser = argparse.ArgumentParser(
        description='Analyze word frequency in documents (.doc, .docx, .pdf)'
    )
    parser.add_argument(
        'file_path',
        help='Path to the document file'
    )
    parser.add_argument(
        '--top', '-t',
        type=int,
        default=100,
        help='Number of top words to display in report (default: 100)'
    )
    parser.add_argument(
        '--output', '-o',
        help='Output file path for the markdown report (default: stdout)'
    )
    parser.add_argument(
        '--language', '-l',
        choices=['chinese', 'english', 'mixed'],
        default='mixed',
        help='Language mode: chinese, english, or mixed (default: mixed)'
    )

    args = parser.parse_args()

    # Validate input file
    if not Path(args.file_path).exists():
        print(f"Error: File not found: {args.file_path}")
        sys.exit(1)

    print(f"Analyzing file: {args.file_path}")

    # Extract text
    print("Extracting text...")
    text = extract_text(args.file_path)

    # Tokenize
    print("Tokenizing text...")
    tokens = tokenize_text(text, language=args.language)

    # Count frequency
    print("Counting word frequency...")
    word_counts = count_word_frequency(tokens)
    total_words = len(tokens)
    unique_words = len(word_counts)

    print(f"Total words: {total_words:,}")
    print(f"Unique words: {unique_words:,}")

    # Generate report
    from datetime import datetime
    timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    print("Generating report...")
    report = generate_markdown_report(
        args.file_path,
        word_counts,
        total_words,
        unique_words,
        top_n=args.top
    )

    # Output report
    if args.output:
        output_path = Path(args.output)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(report, encoding='utf-8')
        print(f"Report saved to: {args.output}")
    else:
        print("\n" + "="*60)
        print(report)


if __name__ == '__main__':
    # For use as a module
    if len(sys.argv) == 1 and __name__ == '__main__':
        # No arguments provided, show help
        print(__doc__)
        sys.exit(0)
    else:
        main()
